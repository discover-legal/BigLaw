"""Render BigLaw's synthesis into the deliverable files a LAB task names.

LAB's judge grades real files, so format matters: .docx via python-docx,
.xlsx via openpyxl (preferring BigLaw's tabulate table when the task produced
one, else markdown tables found in the synthesis), .pdf via PyMuPDF, and
plain text/markdown as-is. `pandoc` on PATH is the fallback for .docx/.pdf
when the Python renderers' libraries are missing.
"""

from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
from pathlib import Path

_TABLE_ROW = re.compile(r"^\s*\|(.+)\|\s*$")
_TABLE_SEP = re.compile(r"^\s*\|?\s*:?-{3,}.*$")
_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET = re.compile(r"^\s*[-*+]\s+(.*)$")
_NUMBERED = re.compile(r"^\s*\d+[.)]\s+(.*)$")
_INLINE_MD = re.compile(r"(\*\*|__|\*|_|`)")


def _strip_inline(text: str) -> str:
    return _INLINE_MD.sub("", text).strip()


def _split_row(line: str) -> list[str]:
    m = _TABLE_ROW.match(line)
    cells = m.group(1).split("|") if m else line.split("|")
    return [_strip_inline(c) for c in cells]


def parse_markdown_tables(text: str) -> list[tuple[list[str], list[list[str]]]]:
    """Extract (header, rows) for every markdown pipe table in text."""
    tables = []
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        if _TABLE_ROW.match(lines[i]) and i + 1 < len(lines) and _TABLE_SEP.match(lines[i + 1]):
            header = _split_row(lines[i])
            rows = []
            i += 2
            while i < len(lines) and _TABLE_ROW.match(lines[i]):
                row = _split_row(lines[i])
                # Pad/trim to header width so openpyxl gets rectangular data.
                row = (row + [""] * len(header))[: len(header)]
                rows.append(row)
                i += 1
            tables.append((header, rows))
        else:
            i += 1
    return tables


# ─── .docx ────────────────────────────────────────────────────────────────────

def _write_docx(text: str, dest: Path) -> None:
    import docx

    doc = docx.Document()
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if _TABLE_ROW.match(line) and i + 1 < len(lines) and _TABLE_SEP.match(lines[i + 1]):
            header = _split_row(line)
            i += 2
            rows = []
            while i < len(lines) and _TABLE_ROW.match(lines[i]):
                rows.append((_split_row(lines[i]) + [""] * len(header))[: len(header)])
                i += 1
            table = doc.add_table(rows=len(rows) + 1, cols=len(header))
            table.style = "Table Grid"
            for c, cell in enumerate(header):
                table.rows[0].cells[c].text = cell
            for r, row in enumerate(rows, 1):
                for c, cell in enumerate(row):
                    table.rows[r].cells[c].text = cell
            continue

        if m := _HEADING.match(line):
            doc.add_heading(_strip_inline(m.group(2)), level=min(len(m.group(1)), 4))
        elif m := _BULLET.match(line):
            doc.add_paragraph(_strip_inline(m.group(1)), style="List Bullet")
        elif m := _NUMBERED.match(line):
            doc.add_paragraph(_strip_inline(m.group(1)), style="List Number")
        elif line.strip():
            doc.add_paragraph(_strip_inline(line))
        i += 1
    doc.save(str(dest))


# ─── .xlsx ────────────────────────────────────────────────────────────────────

def _write_xlsx(text: str, dest: Path, biglaw_table: dict | None) -> None:
    import openpyxl

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    if biglaw_table and biglaw_table.get("columns"):
        cols = [c for c in biglaw_table["columns"] if not c.startswith("_")]
        ws = wb.create_sheet("Results")
        ws.append(cols)
        for row in biglaw_table.get("rows") or []:
            ws.append([row.get(c, "") for c in cols])
    else:
        for n, (header, rows) in enumerate(parse_markdown_tables(text), 1):
            ws = wb.create_sheet(f"Table {n}" if n > 1 else "Results")
            ws.append(header)
            for row in rows:
                ws.append(row)

    if not wb.sheetnames:
        # No tabular content found — preserve the synthesis line-by-line so
        # the judge still sees the substance rather than an empty workbook.
        ws = wb.create_sheet("Results")
        for line in text.splitlines():
            ws.append([line])
    wb.save(str(dest))


# ─── .pdf ─────────────────────────────────────────────────────────────────────

def _write_pdf(text: str, dest: Path) -> None:
    import fitz  # PyMuPDF

    plain = "\n".join(_strip_inline(l) if l.strip() else "" for l in text.splitlines())
    doc = fitz.open()
    rect = fitz.Rect(54, 54, 541, 788)  # letter with 0.75" margins
    lines = plain.splitlines() or [""]
    while lines:
        page = doc.new_page()
        # insert_textbox inserts nothing and returns a negative deficit when
        # the text overflows, so shrinking and retrying on the same page is safe.
        n = len(lines)
        while page.insert_textbox(rect, "\n".join(lines[:n]), fontsize=10, fontname="helv") < 0 and n > 1:
            n = max(1, int(n * 0.7))
        lines = lines[n:]
    doc.save(str(dest))
    doc.close()


# ─── pandoc fallback ──────────────────────────────────────────────────────────

def _pandoc_render(text: str, dest: Path) -> None:
    if not shutil.which("pandoc"):
        raise RuntimeError(f"cannot render {dest.suffix}: renderer library missing and pandoc not on PATH")
    with tempfile.NamedTemporaryFile("w", suffix=".md", delete=False) as tmp:
        tmp.write(text)
        src = tmp.name
    proc = subprocess.run(["pandoc", src, "-o", str(dest)], capture_output=True, text=True, timeout=120)
    Path(src).unlink(missing_ok=True)
    if proc.returncode != 0:
        raise RuntimeError(f"pandoc failed: {proc.stderr.strip()[:200]}")


# ─── Per-deliverable splitting ────────────────────────────────────────────────

_MARKER = re.compile(r"^\s*={2,}\s*DELIVERABLE:\s*`?(.+?)`?\s*={2,}\s*$", re.MULTILINE | re.IGNORECASE)


def split_by_markers(text: str, deliverables: list[str]) -> dict[str, str]:
    """Map deliverable filenames to their `=== DELIVERABLE: name ===` sections.

    Matching is by full relative name or basename, case-insensitive. Returns
    only the deliverables that were found; callers fall back to the full
    synthesis for any that are missing.
    """
    sections: dict[str, str] = {}
    matches = list(_MARKER.finditer(text))
    for i, m in enumerate(matches):
        name = m.group(1).strip().lower()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[m.end():end].strip()
        for d in deliverables:
            if d not in sections and name in (d.lower(), Path(d).name.lower()):
                sections[d] = body
                break
    return sections


def strip_markers(text: str) -> str:
    return _MARKER.sub("", text).strip()


# ─── Entry point ──────────────────────────────────────────────────────────────

def render_deliverable(text: str, dest: Path, biglaw_table: dict | None = None) -> None:
    """Write the synthesis to dest in the format its extension implies."""
    dest.parent.mkdir(parents=True, exist_ok=True)
    ext = dest.suffix.lower()
    try:
        if ext == ".docx":
            _write_docx(text, dest)
        elif ext in (".xlsx", ".xlsm"):
            _write_xlsx(text, dest, biglaw_table)
        elif ext == ".pdf":
            _write_pdf(text, dest)
        else:
            dest.write_text(text, encoding="utf-8")
    except ImportError:
        _pandoc_render(text, dest)
